from __future__ import annotations

from typing import Optional, Annotated, Literal
from datetime import datetime, timezone
from loguru import logger
import json
import io

from langchain_core.tools import tool
from langgraph.prebuilt import InjectedState, create_react_agent

from app.integrations.gigachat_client import gigachat_client
from app.integrations import salutespeech_service
from app.integrations.minio_service import MinioService


class ContentGenerationTools:
    
    def __init__(self, minio_service: MinioService):
        self.minio_service = minio_service
    
    @tool
    async def generate_image(
        self,
        state: Annotated[dict, InjectedState],
        prompt: str,
        width: int = 1024,
        height: int = 1024,
        folder: Optional[str] = None,
    ) -> str:
        """Сгенерировать изображение через GigaChat и сохранить в MinIO.
        
        Используй для:
        - Создания иллюстраций для статей
        - Визуализации концепций
        - Генерации обучающих материалов
        
        Args:
            state: Состояние графа (автоматически инжектится)
            prompt: Описание изображения (детальное, на русском)
            width: Ширина изображения (по умолчанию 1024)
            height: Высота изображения (по умолчанию 1024)
            folder: Папка в MinIO (по умолчанию "generated/images")
        
        Returns:
            JSON с информацией о сгенерированном изображении
        """
        try:
            user_id = state["user_id"]
            
            # Генерируем изображение через GigaChat
            file_id = await gigachat_client.generate_image(
                prompt=prompt,
                width=width,
                height=height
            )
            
            # Скачиваем изображение из GigaChat
            image_bytes = await gigachat_client.download_file(file_id)
            image_io = io.BytesIO(image_bytes)
            
            # Определяем папку
            upload_folder = folder or "generated/images"
            
            # Формируем имя файла
            filename = f"image_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            
            # Сохраняем в MinIO
            minio_object_name = await self.minio_service.upload_file(
                file=image_io,
                filename=filename,
                content_type="image/png",
                folder=upload_folder
            )
            
            # Получаем URL
            minio_url = await self.minio_service.get_file_url(minio_object_name)
            
            # Формируем результат
            result = {
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "prompt": prompt,
                "width": width,
                "height": height,
                "minio_object_name": minio_object_name,
                "minio_url": minio_url,
                "file_size_bytes": len(image_bytes)
            }
            
            # ВАЖНО: Добавляем в state["generated_files"]
            generated_files = state.get("generated_files", [])
            generated_files.append({
                "type": "image",
                "filename": filename,
                "minio_object_name": minio_object_name,
                "minio_url": minio_url,
                "file_size_bytes": len(image_bytes),
                "metadata": {
                    "prompt": prompt,
                    "width": width,
                    "height": height,
                }
            })
            state["generated_files"] = generated_files
            
            logger.info(f"Image generated and saved: {minio_object_name}")
            return json.dumps(result, ensure_ascii=False, indent=2)
            
        except Exception as e:
            logger.error(f"Failed to generate image: {e}")
            return json.dumps({
                "error": str(e),
                "prompt": prompt
            }, ensure_ascii=False)
    
    @tool
    async def create_chart(
        self,
        state: Annotated[dict, InjectedState],
        chart_type: Literal["line", "bar", "pie", "scatter", "table"],
        data: str,
        title: str = "",
        x_label: str = "",
        y_label: str = "",
        folder: Optional[str] = None,
    ) -> str:
        """Создать график, диаграмму или таблицу и сохранить в MinIO.
        
        Используй для:
        - Визуализации данных о здоровье
        - Графиков веса, температуры
        - Диаграмм статистики
        - Таблиц с данными
        
        Args:
            state: Состояние графа (автоматически инжектится)
            chart_type: Тип графика - line/bar/pie/scatter/table
            data: Данные в JSON формате
                  Для line/bar/scatter: {"x": [1,2,3], "y": [4,5,6]} или {"labels": [...], "values": [...]}
                  Для pie: {"labels": ["A", "B"], "values": [30, 70]}
                  Для table: {"columns": ["Col1", "Col2"], "data": [[1,2], [3,4]]}
            title: Заголовок графика
            x_label: Подпись оси X
            y_label: Подпись оси Y
            folder: Папка в MinIO (по умолчанию "generated/charts")
        
        Returns:
            JSON с информацией о созданном графике
        """
        try:
            user_id = state["user_id"]
            
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            
            # Парсим данные
            data_dict = json.loads(data)
            
            # Создаём фигуру
            fig, ax = plt.subplots(figsize=(10, 6))
            
            # Устанавливаем заголовок
            if title:
                ax.set_title(title, fontsize=14, fontweight='bold')
            
            # Создаём график в зависимости от типа
            if chart_type == "line":
                x_data = data_dict.get("x", data_dict.get("labels", []))
                y_data = data_dict.get("y", data_dict.get("values", []))
                ax.plot(x_data, y_data, marker='o', linewidth=2, markersize=6)
                ax.grid(True, alpha=0.3)
                
            elif chart_type == "bar":
                x_data = data_dict.get("x", data_dict.get("labels", []))
                y_data = data_dict.get("y", data_dict.get("values", []))
                ax.bar(x_data, y_data, alpha=0.7, color='#4CAF50')
                ax.grid(True, axis='y', alpha=0.3)
                
            elif chart_type == "pie":
                labels = data_dict.get("labels", [])
                values = data_dict.get("values", [])
                colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#FFA07A', '#98D8C8']
                ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
                ax.axis('equal')
                
            elif chart_type == "scatter":
                x_data = data_dict.get("x", [])
                y_data = data_dict.get("y", [])
                ax.scatter(x_data, y_data, alpha=0.6, s=100, color='#FF6B6B')
                ax.grid(True, alpha=0.3)
                
            elif chart_type == "table":
                ax.axis('tight')
                ax.axis('off')
                
                columns = data_dict.get("columns", [])
                table_data = data_dict.get("data", [])
                
                table = ax.table(
                    cellText=table_data,
                    colLabels=columns,
                    cellLoc='center',
                    loc='center',
                    colWidths=[0.2] * len(columns)
                )
                table.auto_set_font_size(False)
                table.set_fontsize(10)
                table.scale(1, 2)
                
                # Стилизация заголовков
                for (row, col), cell in table.get_celld().items():
                    if row == 0:
                        cell.set_facecolor('#4CAF50')
                        cell.set_text_props(weight='bold', color='white')
                    else:
                        cell.set_facecolor('#F0F0F0' if row % 2 == 0 else 'white')
            
            # Подписи осей (если не таблица и не круговая)
            if chart_type not in ["table", "pie"]:
                if x_label:
                    ax.set_xlabel(x_label, fontsize=11)
                if y_label:
                    ax.set_ylabel(y_label, fontsize=11)
            
            # Сохраняем в буфер
            buffer = io.BytesIO()
            plt.tight_layout()
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
            plt.close(fig)
            
            buffer.seek(0)
            
            # Определяем папку
            upload_folder = folder or "generated/charts"
            
            # Формируем имя файла
            filename = f"chart_{chart_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            
            # Сохраняем в MinIO
            minio_object_name = await self.minio_service.upload_file(
                file=buffer,
                filename=filename,
                content_type="image/png",
                folder=upload_folder
            )
            
            # Получаем URL
            minio_url = await self.minio_service.get_file_url(minio_object_name)
            
            # Формируем результат
            result = {
                "created_at": datetime.now(timezone.utc).isoformat(),
                "chart_type": chart_type,
                "title": title,
                "minio_object_name": minio_object_name,
                "minio_url": minio_url,
                "file_size_bytes": len(buffer.getvalue())
            }
            
            # ВАЖНО: Добавляем в state["generated_files"]
            generated_files = state.get("generated_files", [])
            generated_files.append({
                "type": "chart",
                "filename": filename,
                "minio_object_name": minio_object_name,
                "minio_url": minio_url,
                "file_size_bytes": len(buffer.getvalue()),
                "metadata": {
                    "chart_type": chart_type,
                    "title": title,
                }
            })
            state["generated_files"] = generated_files
            
            logger.info(f"Chart created and saved: {minio_object_name}")
            return json.dumps(result, ensure_ascii=False, indent=2)
            
        except Exception as e:
            logger.error(f"Failed to create chart: {e}")
            return json.dumps({
                "error": str(e),
                "chart_type": chart_type
            }, ensure_ascii=False)
    
    @tool
    async def text_to_speech(
        self,
        state: Annotated[dict, InjectedState],
        text: str,
        voice: str = "Bys_24000",
        audio_format: str = "wav16",
        folder: Optional[str] = None,
    ) -> str:
        """Синтезировать речь из текста и сохранить в MinIO.
        
        Используй для:
        - Озвучивания текстовых ответов
        - Создания аудио-инструкций
        - Голосовых напоминаний
        
        Args:
            state: Состояние графа (автоматически инжектится)
            text: Текст для синтеза речи
            voice: Голос (Bys_24000, Nec_24000, May_24000, Ost_24000, Pon_24000)
            audio_format: Формат аудио (wav16, pcm16, opus)
            folder: Папка в MinIO (по умолчанию "generated/audio")
        
        Returns:
            JSON с информацией о синтезированном аудио
        """
        try:
            user_id = state["user_id"]
            
            # Синтезируем речь через SaluteSpeech
            audio_bytes = await salutespeech_service.text_to_speech(
                text=text,
                voice=voice,
                format=audio_format
            )
            
            audio_io = io.BytesIO(audio_bytes)
            
            # Определяем папку
            upload_folder = folder or "generated/audio"
            
            # Определяем расширение файла
            extension_map = {
                "wav16": "wav",
                "pcm16": "pcm",
                "opus": "opus"
            }
            ext = extension_map.get(audio_format, "wav")
            
            # Формируем имя файла
            filename = f"tts_{voice}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"
            
            # Определяем content-type
            content_type_map = {
                "wav16": "audio/wav",
                "pcm16": "audio/pcm",
                "opus": "audio/opus"
            }
            content_type = content_type_map.get(audio_format, "audio/wav")
            
            # Сохраняем в MinIO
            minio_object_name = await self.minio_service.upload_file(
                file=audio_io,
                filename=filename,
                content_type=content_type,
                folder=upload_folder
            )
            
            # Получаем URL
            minio_url = await self.minio_service.get_file_url(minio_object_name)
            
            # Формируем результат
            result = {
                "synthesized_at": datetime.now(timezone.utc).isoformat(),
                "text_preview": text[:100] + ("..." if len(text) > 100 else ""),
                "text_length": len(text),
                "voice": voice,
                "format": audio_format,
                "minio_object_name": minio_object_name,
                "minio_url": minio_url,
                "file_size_bytes": len(audio_bytes)
            }
            
            # ВАЖНО: Добавляем в state["generated_files"]
            generated_files = state.get("generated_files", [])
            generated_files.append({
                "type": "audio",
                "filename": filename,
                "minio_object_name": minio_object_name,
                "minio_url": minio_url,
                "file_size_bytes": len(audio_bytes),
                "metadata": {
                    "voice": voice,
                    "format": audio_format,
                    "text_preview": text[:100] + ("..." if len(text) > 100 else ""),
                }
            })
            state["generated_files"] = generated_files
            
            logger.info(f"TTS generated and saved: {minio_object_name}")
            return json.dumps(result, ensure_ascii=False, indent=2)
            
        except Exception as e:
            logger.error(f"Failed to synthesize speech: {e}")
            return json.dumps({
                "error": str(e),
                "text_preview": text[:50]
            }, ensure_ascii=False)
    
    @tool
    async def generate_pdf_report(
        self,
        state: Annotated[dict, InjectedState],
        title: str,
        content: str,
        folder: Optional[str] = None,
    ) -> str:
        """Создать PDF отчёт и сохранить в MinIO.
        
        Используй для:
        - Отчётов о здоровье питомца
        - Медицинских справок
        - Сводок по питанию
        
        Args:
            state: Состояние графа (автоматически инжектится)
            title: Заголовок отчёта
            content: Содержимое отчёта (поддерживает простую разметку: **жирный**)
            folder: Папка в MinIO (по умолчанию "generated/reports")
        
        Returns:
            JSON с информацией о созданном PDF
        """
        try:
            user_id = state["user_id"]
            
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Регистрируем шрифт с поддержкой кириллицы
            try:
                pdfmetrics.registerFont(TTFont('DejaVuSans', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
                pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
                font_name = 'DejaVuSans'
                font_name_bold = 'DejaVuSans-Bold'
            except:
                logger.warning("DejaVu fonts not found, using default")
                font_name = 'Helvetica'
                font_name_bold = 'Helvetica-Bold'
            
            # Создаём PDF в памяти
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            
            # Стили
            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=font_name_bold,
                fontSize=18,
                textColor='#2C3E50',
                spaceAfter=20,
                alignment=1  # Center
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontName=font_name,
                fontSize=11,
                leading=16,
                spaceAfter=12,
            )
            
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=9,
                textColor='#888888',
            )
            
            # Формируем содержимое
            story = []
            
            # Заголовок
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.5*cm))
            
            # Дата создания
            date_text = f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            story.append(Paragraph(date_text, footer_style))
            story.append(Spacer(1, 0.8*cm))
            
            # Основной контент (разбиваем по параграфам)
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Простая обработка **жирный**
                    para_text = para.replace('**', '<b>').replace('**', '</b>')
                    story.append(Paragraph(para_text, body_style))
            
            # Генерируем PDF
            doc.build(story)
            
            buffer.seek(0)
            
            # Определяем папку
            upload_folder = folder or "generated/reports"
            
            # Формируем имя файла
            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            # Сохраняем в MinIO
            minio_object_name = await self.minio_service.upload_file(
                file=buffer,
                filename=filename,
                content_type="application/pdf",
                folder=upload_folder
            )
            
            # Получаем URL
            minio_url = await self.minio_service.get_file_url(minio_object_name)
            
            # Формируем результат
            result = {
                "created_at": datetime.now(timezone.utc).isoformat(),
                "title": title,
                "content_length": len(content),
                "minio_object_name": minio_object_name,
                "minio_url": minio_url,
                "file_size_bytes": len(buffer.getvalue())
            }
            
            # ВАЖНО: Добавляем в state["generated_files"]
            generated_files = state.get("generated_files", [])
            generated_files.append({
                "type": "pdf",
                "filename": filename,
                "minio_object_name": minio_object_name,
                "minio_url": minio_url,
                "file_size_bytes": len(buffer.getvalue()),
                "metadata": {
                    "title": title,
                }
            })
            state["generated_files"] = generated_files
            
            logger.info(f"PDF report created and saved: {minio_object_name}")
            return json.dumps(result, ensure_ascii=False, indent=2)
            
        except Exception as e:
            logger.error(f"Failed to generate PDF report: {e}")
            return json.dumps({
                "error": str(e),
                "title": title
            }, ensure_ascii=False)
    
    @tool
    async def generate_docx_report(
        self,
        state: Annotated[dict, InjectedState],
        title: str,
        content: str,
        folder: Optional[str] = None,
    ) -> str:
        """Создать DOCX отчёт и сохранить в MinIO.
        
        Используй для:
        - Редактируемых отчётов
        - Документов для печати
        - Шаблонов для заполнения
        
        Args:
            state: Состояние графа (автоматически инжектится)
            title: Заголовок отчёта
            content: Содержимое отчёта (поддерживает **жирный** текст)
            folder: Папка в MinIO (по умолчанию "generated/reports")
        
        Returns:
            JSON с информацией о созданном DOCX
        """
        try:
            user_id = state["user_id"]
            
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Создаём документ
            doc = Document()
            
            # Заголовок
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Дата
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(128, 128, 128)
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_paragraph()  # Пустая строка
            
            # Основной контент
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()
                    
                    # Простая обработка **жирный**
                    parts = para_text.split('**')
                    for i, part in enumerate(parts):
                        if part:
                            run = para.add_run(part)
                            if i % 2 == 1:  # Нечётные части - жирные
                                run.bold = True
                            run.font.size = Pt(11)
            
            # Сохраняем в буфер
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Определяем папку
            upload_folder = folder or "generated/reports"
            
            # Формируем имя файла
            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            # Сохраняем в MinIO
            minio_object_name = await self.minio_service.upload_file(
                file=buffer,
                filename=filename,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                folder=upload_folder
            )
            
            # Получаем URL
            minio_url = await self.minio_service.get_file_url(minio_object_name)
            
            # Формируем результат
            result = {
                "created_at": datetime.now(timezone.utc).isoformat(),
                "title": title,
                "content_length": len(content),
                "minio_object_name": minio_object_name,
                "minio_url": minio_url,
                "file_size_bytes": len(buffer.getvalue())
            }
            
            # ВАЖНО: Добавляем в state["generated_files"]
            generated_files = state.get("generated_files", [])
            generated_files.append({
                "type": "docx",
                "filename": filename,
                "minio_object_name": minio_object_name,
                "minio_url": minio_url,
                "file_size_bytes": len(buffer.getvalue()),
                "metadata": {
                    "title": title,
                }
            })
            state["generated_files"] = generated_files
            
            logger.info(f"DOCX report created and saved: {minio_object_name}")
            return json.dumps(result, ensure_ascii=False, indent=2)
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX report: {e}")
            return json.dumps({
                "error": str(e),
                "title": title
            }, ensure_ascii=False)


def create_content_generation_agent(
    minio_service: MinioService,
    llm,
    name: str = "content_generation",
):
    """Создать агента для генерации контента (изображения, графики, аудио, отчёты)
    
    Args:
        minio_service: Сервис для сохранения файлов в MinIO
        llm: Языковая модель
        name: Имя агента (для supervisor handoff)
    
    Returns:
        Compiled ReAct agent
    """
    tools_instance = ContentGenerationTools(minio_service)
    
    tools = [
        tools_instance.generate_image,
        tools_instance.create_chart,
        tools_instance.text_to_speech,
        tools_instance.generate_pdf_report,
        tools_instance.generate_docx_report,
    ]
    
    prompt = (
        "Ты - эксперт по генерации контента.\n\n"
        "Твои возможности:\n"
        "- Генерация изображений (GigaChat)\n"
        "- Создание графиков и таблиц (matplotlib)\n"
        "- Синтез речи (SaluteSpeech TTS)\n"
        "- Генерация отчётов (PDF, DOCX)\n\n"
        "ВСЕ сгенерированные файлы АВТОМАТИЧЕСКИ сохраняются в MinIO и добавляются в state.\n\n"
        "Важно:\n"
        "- Для generate_image используй детальные русские промпты\n"
        "- Для create_chart добавляй title, x_label, y_label для читаемости\n"
        "- Для text_to_speech передавай весь текст целиком в параметр text\n"
        "- Каждый инструмент возвращает minio_url для доступа к файлу\n\n"
        "Все tools возвращают JSON - просто верни результат как есть."
    )
    
    agent = create_react_agent(
        model=llm,
        tools=tools,
        name=name,
        prompt=prompt,
    )
    
    logger.info(f"Created ContentGenerationAgent '{name}' with {len(tools)} tools")
    return agent